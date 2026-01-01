import cv2
import sys
import os
from docx import Document
from docx.shared import Inches
import io

def extract_frames_to_word(video_path, output_docx='video_frames.docx'):
    """
    Extract one frame per second from a video and insert them into a Word document
    with timestamps.
    """
    # Check if video file exists
    if not os.path.exists(video_path):
        print(f"Error: Video file '{video_path}' not found.")
        sys.exit(1)
    
    # Open video file
    cap = cv2.VideoCapture(video_path)
    
    if not cap.isOpened():
        print(f"Error: Could not open video file '{video_path}'.")
        sys.exit(1)
    
    # Get video properties
    fps = cap.get(cv2.CAP_PROP_FPS)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    duration = total_frames / fps if fps > 0 else 0
    
    print(f"Video FPS: {fps:.2f}")
    print(f"Total frames: {total_frames}")
    print(f"Duration: {duration:.2f} seconds")
    print("Extracting frames...")
    
    # Create Word document
    doc = Document()
    doc.add_heading('Video Frames', 0)
    
    frame_count = 0
    extracted_count = 0
    
    # Extract one frame per second
    while True:
        ret, frame = cap.read()
        
        if not ret:
            break
        
        # Calculate current time in seconds
        current_time = frame_count / fps if fps > 0 else 0
        
        # Extract frame at each second (or close to it)
        if frame_count % int(fps) == 0 or frame_count == 0:
            # Convert timestamp to MM:SS format
            minutes = int(current_time // 60)
            seconds = int(current_time % 60)
            timestamp = f"{minutes:02d}:{seconds:02d}"
            
            # Add timestamp as paragraph
            doc.add_paragraph(f"Time: {timestamp}")
            
            # Convert frame to image bytes
            success, buffer = cv2.imencode('.png', frame)
            if success:
                image_bytes = io.BytesIO(buffer)
                
                # Add image to document (width of 6 inches to fit page)
                doc.add_picture(image_bytes, width=Inches(6))
            
            extracted_count += 1
            print(f"Extracted frame at {timestamp} ({extracted_count} total)")
        
        frame_count += 1
    
    # Release video capture
    cap.release()
    
    # Save document
    doc.save(output_docx)
    print(f"\nDocument saved as '{output_docx}'")
    print(f"Total frames extracted: {extracted_count}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python script.py <video_file_path>")
        print("Example: python script.py video.mp4")
        sys.exit(1)
    
    video_path = sys.argv[1]
    extract_frames_to_word(video_path)

